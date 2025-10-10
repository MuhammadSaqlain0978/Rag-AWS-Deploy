import PyPDF2
import json
import os
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from mistralai import Mistral
from docx import Document as DocxDocument
import docx2txt
import mammoth
from combined_script import combine_dataset_files

# from combined1 import combine_dataset_files

class UniversityRAGBot:
    """Enhanced RAG-based chatbot for university queries with database folder support"""
    
    def __init__(self, api_key, model_name="open-mistral-7b"):
        self.mistral_client = Mistral(api_key=api_key)
        self.model_name = model_name
        self.vector_store = None
        self.embeddings = None
        self.supported_extensions = {'.pdf', '.docx', '.doc', '.txt', '.json'}
        print(f"🤖 Using model: {model_name}")

    


        
    def load_university_data(self, database_folder="comfolder", vector_store_path="./vector_store"):
        """Load and process all documents from database folder and save/load vector store"""
        print(f"📚 Loading university data from database folder: {database_folder}")
        
        # Initialize embeddings
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        
        # Check if vector store already exists
        if os.path.exists(vector_store_path) and os.path.exists(os.path.join(vector_store_path, "index.faiss")):
            print("🔄 Loading existing vector store...")
            try:
                self.vector_store = FAISS.load_local(vector_store_path, self.embeddings, allow_dangerous_deserialization=True)
                print("✅ Loaded existing vector store!")
                return
            except Exception as e:
                print(f"⚠️ Error loading vector store: {e}. Rebuilding vector store...")
        
        # Check if database folder exists
        if not os.path.exists(database_folder):
            raise ValueError(f"Database folder '{database_folder}' not found! Please create the folder and add your documents.")
        
        # Load all documents from database folder
        all_documents = self._load_database_folder(database_folder)
        
        if not all_documents:
            raise ValueError(f"No supported documents found in '{database_folder}' folder! Please add PDF, Word, or text files.")
        
        # Split documents into chunks for better retrieval
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,  # Smaller chunks for university info
            chunk_overlap=200,   # Some overlap to maintain context
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", "", "\n========================================================================================"]
        )
        
        # Split all documents
        split_docs = text_splitter.split_documents(all_documents)
        print(f"✅ Created {len(split_docs)} knowledge chunks from all documents")
        
        # Build vector database from documents
        self.vector_store = FAISS.from_documents(split_docs, self.embeddings)
        
        # Save vector store to disk
        try:
            if not os.path.exists(vector_store_path):
                os.makedirs(vector_store_path)
            self.vector_store.save_local(vector_store_path)
            print(f"✅ Vector store saved to {vector_store_path}")
        except Exception as e:
            print(f"⚠️ Error saving vector store: {e}")
        
        print("✅ University knowledge base ready with all database documents!")
        
    def _load_database_folder(self, database_folder):
        """Load all supported documents from the database folder"""
        documents = []
        file_stats = {
            'pdf': 0,
            'docx': 0, 
            'doc': 0,
            'txt': 0,
            'json': 0,
            'failed': 0
        }
        
        print(f"🔍 Scanning database folder: {database_folder}")
        
        # Walk through all files in the database folder (including subfolders)
        for root, dirs, files in os.walk(database_folder):
            for filename in files:
                file_path = os.path.join(root, filename)
                file_ext = os.path.splitext(filename)[1].lower()
                
                # Skip unsupported file types
                if file_ext not in self.supported_extensions:
                    continue
                
                print(f"📄 Processing: {filename}")
                
                try:
                    # Process based on file type
                    if file_ext == '.pdf':
                        docs = self._load_pdf_file(file_path, filename)
                        file_stats['pdf'] += 1
                        
                    elif file_ext == '.docx':
                        docs = self._load_docx_file(file_path, filename)
                        file_stats['docx'] += 1
                        
                    elif file_ext == '.doc':
                        docs = self._load_doc_file(file_path, filename)
                        file_stats['doc'] += 1
                        
                    elif file_ext == '.txt':
                        docs = self._load_txt_file(file_path, filename)
                        file_stats['txt'] += 1
                        
                    elif file_ext == '.json':
                        docs = self._load_json_file(file_path, filename)
                        file_stats['json'] += 1
                    
                    documents.extend(docs)
                    print(f"✅ Successfully loaded: {filename} ({len(docs)} chunks)")
                    
                except Exception as e:
                    print(f"❌ Error loading {filename}: {e}")
                    file_stats['failed'] += 1
        
        # Print summary
        print(f"\n📊 Loading Summary:")
        print(f"   📄 PDF files: {file_stats['pdf']}")
        print(f"   📝 Word files (.docx): {file_stats['docx']}")
        print(f"   📝 Word files (.doc): {file_stats['doc']}")
        print(f"   📃 Text files: {file_stats['txt']}")
        print(f"   🌐 JSON files: {file_stats['json']}")
        print(f"   ❌ Failed files: {file_stats['failed']}")
        print(f"   📚 Total documents: {len(documents)}")
        
        return documents
    
    def _load_pdf_file(self, file_path, filename):
        """Load a single PDF file"""
        documents = []
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # Extract text from each page
                full_text = ""
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        full_text += page_text + "\n\n"
                
                if full_text.strip():
                    doc = Document(
                        page_content=full_text,
                        metadata={
                            'source': filename,
                            'source_type': 'pdf',
                            'file_path': file_path,
                            'pages': len(reader.pages),
                            'category': 'university_document'
                        }
                    )
                    documents.append(doc)
                    
        except Exception as e:
            raise Exception(f"Error reading PDF {filename}: {e}")
        
        return documents
    
    def _load_docx_file(self, file_path, filename):
        """Load a single DOCX file"""
        documents = []
        try:
            # Method 1: Using docx2txt (simpler but less metadata)
            text = docx2txt.process(file_path)
            
            if text and text.strip():
                doc = Document(
                    page_content=text,
                    metadata={
                        'source': filename,
                        'source_type': 'docx',
                        'file_path': file_path,
                        'category': 'university_document'
                    }
                )
                documents.append(doc)
            else:
                # Method 2: Using python-docx (more detailed)
                docx_doc = DocxDocument(file_path)
                paragraphs = []
                for paragraph in docx_doc.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text)
                
                if paragraphs:
                    text = '\n\n'.join(paragraphs)
                    doc = Document(
                        page_content=text,
                        metadata={
                            'source': filename,
                            'source_type': 'docx',
                            'file_path': file_path,
                            'paragraphs': len(paragraphs),
                            'category': 'university_document'
                        }
                    )
                    documents.append(doc)
                    
        except Exception as e:
            raise Exception(f"Error reading DOCX {filename}: {e}")
        
        return documents
    
    def _load_doc_file(self, file_path, filename):
        """Load a single DOC file ( older Word format )"""
        documents = []
        try:
            # For .doc files, we'll try to use mammoth if available
            # Otherwise, suggest converting to .docx
            try:
                with open(file_path, "rb") as docx_file:
                    result = mammoth.extract_raw_text(docx_file)
                    text = result.value
                    
                if text and text.strip():
                    doc = Document(
                        page_content=text,
                        metadata={
                            'source': filename,
                            'source_type': 'doc',
                            'file_path': file_path,
                            'category': 'university_document',
                            'note': 'Converted from .doc format'
                        }
                    )
                    documents.append(doc)
            except ImportError:
                print(f"⚠️ Cannot read .doc file {filename}. Please install 'mammoth' or convert to .docx format")
                raise Exception("mammoth library not available for .doc files")
                
        except Exception as e:
            raise Exception(f"Error reading DOC {filename}: {e}")
        
        return documents
    
    def _load_txt_file(self, file_path, filename):
        """Load a single text file"""
        documents = []
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            if content.strip():
                doc = Document(
                    page_content=content,
                    metadata={
                        'source': filename,
                        'source_type': 'txt',
                        'file_path': file_path,
                        'size': len(content),
                        'category': 'university_document'
                    }
                )
                documents.append(doc)
                
        except UnicodeDecodeError:
            # Try different encodings
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                if content.strip():
                    doc = Document(
                        page_content=content,
                        metadata={
                            'source': filename,
                            'source_type': 'txt',
                            'file_path': file_path,
                            'encoding': 'latin-1',
                            'category': 'university_document'
                        }
                    )
                    documents.append(doc)
            except Exception as e:
                raise Exception(f"Error reading text file with different encodings: {e}")
        except Exception as e:
            raise Exception(f"Error reading TXT {filename}: {e}")
        
        return documents
    
    def _load_json_file(self, file_path, filename):
        """Load a single JSON file"""
        documents = []
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Handle different JSON structures
            if isinstance(data, list):
                # Array of objects
                for i, item in enumerate(data):
                    if isinstance(item, dict):
                        content = item.get('content') or item.get('text') or str(item)
                        if content and len(str(content).strip()) > 20:
                            doc = Document(
                                page_content=str(content),
                                metadata={
                                    'source': f"{filename}[{i}]",
                                    'source_type': 'json',
                                    'file_path': file_path,
                                    'item_id': item.get('id', i),
                                    'category': item.get('category', 'university_document')
                                }
                            )
                            documents.append(doc)
                            
            elif isinstance(data, dict):
                # Single object or nested structure
                if 'content' in data or 'text' in data:
                    content = data.get('content') or data.get('text')
                    if content and len(str(content).strip()) > 20:
                        doc = Document(
                            page_content=str(content),
                            metadata={
                                'source': filename,
                                'source_type': 'json',
                                'file_path': file_path,
                                'category': data.get('category', 'university_document')
                            }
                        )
                        documents.append(doc)
                else:
                    # Treat entire dict as content
                    content = json.dumps(data, indent=2)
                    doc = Document(
                        page_content=content,
                        metadata={
                            'source': filename,
                            'source_type': 'json',
                            'file_path': file_path,
                            'category': 'university_document'
                        }
                    )
                    documents.append(doc)
                    
        except Exception as e:
            raise Exception(f"Error reading JSON {filename}: {e}")
        
        return documents
    
    def answer_query(self, user_question,  chat_history=None):
        """Answer university-related queries using RAG with enhanced context"""
        if not self.vector_store:
            return "❌ Please load university data first!"
            
        try:
            # Step 1: Find relevant information from university dataset
            relevant_docs = self.vector_store.similarity_search(
                user_question, 
                k=3  # Get top 6 most relevant chunks
            )
            
            # Step 2: Organize context by source type
            context_by_type = {
                'pdf': [],
                'docx': [],
                'doc': [],
                'txt': [],
                'json': []
            }
            
            for doc in relevant_docs:
                source_type = doc.metadata.get('source_type', 'unknown')
                content = doc.page_content.strip()
                source = doc.metadata.get('source', 'Unknown')
                
                if source_type in context_by_type:
                    context_by_type[source_type].append(f"[{source}] {content}")
            
            # Step 3: Build structured context
            context_parts = []
            
            for doc_type, contents in context_by_type.items():
                if contents:
                    type_name = {
                        'pdf': 'PDF DOCUMENTS',
                        'docx': 'WORD DOCUMENTS (.docx)',
                        'doc': 'WORD DOCUMENTS (.doc)', 
                        'txt': 'TEXT FILES',
                        'json': 'JSON DATA'
                    }.get(doc_type, doc_type.upper())
                    
                    context_parts.append(f"{type_name}:\n" + "\n\n".join(contents))
            
            full_context = "\n\n" + "="*50 + "\n\n".join(context_parts)
            
            # Step 4: Create enhanced prompt
            prompt = f"""You are a helpful University assistant. Answer the student's question based ONLY on the provided university information below.

Guidelines for your response:
- Be helpful, accurate, and student-friendly
- Provide complete and detailed answers when information is available
- If specific information is not available, politely say you don't have that information
- Always be encouraging and supportive in your tone
- Focus on the most relevant information for the student's question
- answer in bullets if the responce requires to be in bullets
- provide the full information

University Database Information:
{full_context}

Student Question: {user_question}

University Assistant Response:"""

            # Step 5: Get response from Mistral
            response = self.mistral_client.chat.complete(
                model=self.model_name,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=700,
                temperature=1.0,
            )
            
            answer = response.choices[0].message.content
            
            # Step 6: Collect source information
            sources_info = []
            source_types_used = set()
            
            for doc in relevant_docs:
                source_info = {
                    'type': doc.metadata.get('source_type', 'unknown'),
                    'source': doc.metadata.get('source', 'Unknown'),
                    'file_path': doc.metadata.get('file_path', ''),
                    'preview': doc.page_content[:100] + "..." if len(doc.page_content) > 100 else doc.page_content
                }
                sources_info.append(source_info)
                source_types_used.add(source_info['type'])
            
            return {
                'answer': answer,
                'sources': sources_info,
                'context_used': len(relevant_docs),
                'source_types': list(source_types_used)
            }
            
        except Exception as e:
            return {
                'answer': f"Sorry, I encountered an error: {str(e)}",
                'sources': [],
                'context_used': 0,
                'source_types': []
            }

def main():
    """Main chatbot interface with database folder support"""
    print("🎓 Enhanced University RAG Chatbot - Database Folder Edition")
    print("=" * 70)
    
    # Get API key and model choice
    api_key = "3XroHID0PrEB4ouyEhBA9R60yzLIV8lf"
    
    print("\n🤖 Choose your model:")
    print("1. mistral-small-latest (Fastest, Cheapest) - Recommended")
    print("2. mistral-large-latest (Best Quality)")
    print("3. open-mistral-7b (Free alternative)")
    
    # model_choice = input("Enter choice (1-3) or press Enter for default (1): ").strip()
    
    # model_map = {
    #     "1": "mistral-small-latest",
    #     "2": "mistral-large-latest", 
    #     "3": "open-mistral-7b",
    #     # "": "mistral-small-latest"
    # }
    
    # selected_model = model_map.get(model_choice)
    selected_model = "open-mistral-7b"


    
    # Initialize chatbot
    bot = UniversityRAGBot(api_key, selected_model)
    
    # Configure database folder
    print(f"\n📁 Database Configuration")
   
    database_folder = "./comfolder"  # Default folder name
    
    # Vector store path
    vector_store_path = "./database_vector_store2"
    
    print(f"\n📋 Supported file types:")
    print("   📄 PDF files (.pdf)")
    print("   📝 Word documents (.docx, .doc)")
    print("   📃 Text files (.txt)")
    print("   🌐 JSON files (.json)")
    
    try:
        bot.load_university_data(
            database_folder=database_folder,
            vector_store_path=vector_store_path
        )
    except Exception as e:
        print(f"❌ Error loading data: {e}")
        print(f"💡 Make sure the '{database_folder}' folder exists and contains supported files.")
        return
    
    print(f"\n🤖 University Assistant is ready!")
    print(f"📚 Knowledge base loaded from: {database_folder}")
    
    print("\n❓ Ask me about:")
    print("   • Academic programs and courses")
    print("   • Admission requirements and procedures")
    print("   • University facilities and services")
    print("   • Contact information and locations")
    print("   • Any information from your uploaded documents")
    print("\n📝 Type 'exit' to quit")
    print("=" * 70)
    
    # Chat loop
    while True:
        question = input("\n🎓 Student: ").strip()
        
        if question.lower() in ['exit', 'quit', 'bye']:
            print("👋 Goodbye! Good luck with your studies!")
            break
            
        if not question:
            print("❓ Please ask a question!")
            continue
            
        # Get answer
        print("🔍 Searching database...")
        result = bot.answer_query(question)
        
        # Display answer
        print(f"\n🤖 Assistant: {result['answer']}")
        
        # Show source info
        if result['context_used'] > 0:
            source_types = ", ".join(result['source_types'])
            print(f"\n📚 (Based on {result['context_used']} sections from: {source_types})")
            
            # Optional detailed sources
            show_sources = input("🔍 Show source files? (y/n): ").lower().strip()
            if show_sources == 'y':
                print("\n📋 Source Files:")
                unique_sources = {}
                for source in result['sources']:
                    key = f"{source['type']}-{source['source']}"
                    if key not in unique_sources:
                        unique_sources[key] = source
                
                for i, (key, source) in enumerate(unique_sources.items(), 1):
                    print(f"{i}. [{source['type'].upper()}] {source['source']}")
                    if source['file_path']:
                        print(f"   📁 Path: {source['file_path']}")
                    print()

if __name__ == "__main__":


    # Step 1: Combine dataset files into one file before building vectors
    try:
        # combine_dataset_files(dataset_folder="dataset", output_file="comfolder/combinedfile2.txt") # from combined1.py
        
        combine_dataset_files(dataset_folder="dataset", output_file="comfolder/combinedfile1.txt") # from combined_script.py
    except Exception as e:
        print(f"❌ Error combining dataset files: {e}")
        
    # Check for required packages
    try:
        import docx2txt
        from docx import Document as DocxDocument
    except ImportError:
        print("❌ Missing required packages!")
        print("Please install: pip install python-docx docx2txt")
        print("For .doc files, also install: pip install mammoth")
        exit(1)
    
    main()